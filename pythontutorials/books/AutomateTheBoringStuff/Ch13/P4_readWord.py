"""Read Word
This program uses :py:mod:`docx` to read Word documents.

"""


def main():
    import docx

    # Reading Word Documents
    doc = docx.Document("demo.docx")
    print(len(doc.paragraphs))
    print(doc.paragraphs[0].text)
    print(doc.paragraphs[1].text)

    print(len(doc.paragraphs[1].runs))
    print(doc.paragraphs[1].runs[0].text)
    print(doc.paragraphs[1].runs[1].text)
    print(doc.paragraphs[1].runs[2].text)
    print(doc.paragraphs[1].runs[3].text)
    print(doc.paragraphs[1].runs[4].text)

    # Getting the Full Text from a .docx File
    # Don't do this, imports should be at the top of the file
    import pythontutorials.books.AutomateTheBoringStuff.Ch13.P5_readDocx as readDocx
    print(readDocx.getText("demo.docx"))


if __name__ == '__main__':
    main()
