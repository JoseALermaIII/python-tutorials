"""Invitations

Say you have a text file of guest names. This guests.txt file has one name
per line, as follows::

    Prof. Plum
    Miss Scarlet
    Col. Mustard
    Al Sweigart
    RoboCop

Write a program that would generate a Word document with custom invitations.

Notes:
    * Uses provided ``guests.txt`` file.
    * Default output file is ``invitations.docx``.
"""


def main():
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Pt

    INTRO = "It would be a pleasure to have the company of"
    ADDRESS = "at 11010 Memory Lane on the Evening of"
    DATE = "April 1st"
    TIME = "at 7 o'clock"

    # Get guest list
    with open("./guests.txt") as file:
        guests = file.read().splitlines()

    # Create styles for invitations
    document = Document()
    styles = document.styles

    # Add Script style
    style = styles.add_style("Script", WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = "MathJax_Caligraphic"  # Only script font in default Ubuntu 18.04
    font.size = Pt(14)
    font.all_caps = True  # Script effect only applies in caps for above font
    font.bold = True
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add Name style
    style = styles.add_style("Name", WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = "Liberation Sans"  # Font in default Ubuntu 18.04
    font.size = Pt(20)
    font.bold = True
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add Date Style
    style = styles.add_style("Date", WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = "Liberation Sans"
    font.size = Pt(14)
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Make each invitation
    for guest in guests:
        document.add_paragraph(INTRO, style="Script")
        document.add_paragraph(guest, style="Name")
        document.add_paragraph(ADDRESS, style="Script")
        document.add_paragraph(DATE, style="Date")
        document.add_paragraph(TIME, style="Script")
        document.add_page_break()

    # Save invitations
    document.save("invitations.docx")


if __name__ == '__main__':
    main()
