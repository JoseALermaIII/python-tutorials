# This program uses the python-docx module to manipulate Word documents
#
# Note:
# - Example .docx files can be downloaded from http://nostarch.com/automatestuff/


import docx

# Changing Run Attributes
doc = docx.Document("demo.docx")
print(doc.paragraphs[0].text)
print(doc.paragraphs[0].style)
doc.paragraphs[0].style = "Normal"
print(doc.paragraphs[1].text)
print((doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text,
       doc.paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text))
doc.paragraphs[1].runs[0].style = "Quote Char"
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True
doc.save("restyled.docx")

# Writing Word Documents
doc = docx.Document()
print(doc.add_paragraph("Hello world!"))
doc.save("helloworld.docx")

doc = docx.Document()
print(doc.add_paragraph("Hello world!"))
paraObj1 = doc.add_paragraph("This is a second paragraph.")
paraObj2 = doc.add_paragraph("This is yet another paragraph.")
print(paraObj1.add_run(" This text is being added to the second paragraph."))
print(doc.add_paragraph("Hello world!", "Title"))
doc.save("multipleParagraphs.docx")

# Adding Headings
doc = docx.Document()
doc.add_heading("Header 0", 0)
doc.add_heading("Header 1", 1)
doc.add_heading("Header 2", 2)
doc.add_heading("Header 3", 3)
doc.add_heading("Header 4", 4)
doc.save("headings.docx")

# Adding Line and Page Breaks
doc = docx.Document()
doc.add_paragraph("This is on the first page!")
doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.add_paragraph("This is on the second page!")
doc.save("twoPage.docx")

# Adding Pictures
doc = docx.Document()
doc.add_picture("zophie.png", width=docx.shared.Inches(1),
                height=docx.shared.Cm(4))
doc.save("picture.docx")
